import streamlit as st
import os
import tempfile
import re
import json
import pandas as pd
from docx import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
import numpy as np
import requests

# Initialize session state
if 'processed' not in st.session_state:
    st.session_state.processed = False
if 'chunks' not in st.session_state:
    st.session_state.chunks = []
if 'vector_store' not in st.session_state:
    st.session_state.vector_store = None
if 'ollama_available' not in st.session_state:
    st.session_state.ollama_available = False

# Subject categories
SUBJECTS = [
    "Operating System", "Computer Network", "DBMS", "OOPS", 
    "System Design", "LLD", "Git", "Linux", "Aptitude"
]

def check_ollama_availability():
    """Check if Ollama is running and has the model"""
    try:
        response = requests.get("http://localhost:11434/api/tags", timeout=5)
        if response.status_code == 200:
            models = response.json().get('models', [])
            # Check if llama3.2 or any model is available
            return len(models) > 0
        return False
    except:
        return False

def clean_text(text):
    """Comprehensive text cleaning with word boundary reconstruction"""
    # Step 1: Fix specific OCR artifacts
    text = re.sub(r'(\b)e([A-Z])', r'\1\2', text)
    text = re.sub(r'\b(\d+);\b', r'\1', text)
    
    # Step 2: Reconstruct word boundaries
    text = re.sub(r'([a-z])([A-Z])', r'\1 \2', text)
    text = re.sub(r'([A-Z])([A-Z][a-z])', r'\1 \2', text)
    text = re.sub(r'([0-9])([A-Za-z])', r'\1 \2', text)
    text = re.sub(r'([A-Za-z])([0-9])', r'\1 \2', text)
    
    # Step 3: Standard cleaning
    text = re.sub(r'[^\w\s\.\-\+\*/\^\(\)\[\]\{\}=,;:&@#%$]', '', text)
    text = re.sub(r'\s+', ' ', text)
    text = re.sub(r'\s([\.\,\;:])', r'\1', text)
    text = re.sub(r'\bpage\s*\d+\b', '', text, flags=re.IGNORECASE)
    text = re.sub(r'(\w)-\s+(\w)', r'\1\2', text)
    
    # Step 4: Fix common OCR errors
    replacements = {
        r'\boperatingsystem\b': 'operating system',
        r'\baddressspace\b': 'address space',
        r'\bmultitasking\b': 'multi-tasking',
        r'\bmainmemory\b': 'main memory',
        r'\bcpubound\b': 'CPU-bound',
        r'\biobound\b': 'I/O-bound',
        r'\bdatastructure\b': 'data structure',
        r'\balgorithm\b': 'algorithm',
    }
    for pattern, replacement in replacements.items():
        text = re.sub(pattern, replacement, text, flags=re.IGNORECASE)
    
    return text.strip()

def extract_pdf_content(pdf_path):
    """Extract text from PDF using alternative methods"""
    text_content = ""
    
    try:
        # Try using pdf2image + pytesseract
        import pdf2image
        import pytesseract
        import cv2
        
        images = pdf2image.convert_from_path(pdf_path)
        
        for i, image in enumerate(images):
            st.info(f"Processing page {i+1}/{len(images)}")
            
            # Convert to grayscale
            gray = cv2.cvtColor(np.array(image), cv2.COLOR_RGB2GRAY)
            
            # Extract text using OCR
            page_text = pytesseract.image_to_string(gray)
            cleaned_page_text = clean_text(page_text)
            text_content += cleaned_page_text + "\n\n"
            
    except Exception as e:
        st.warning(f"OCR method failed: {e}. ")
        # Fallback to PyPDF2
        # try:
        #     import PyPDF2
        #     with open(pdf_path, 'rb') as file:
        #         reader = PyPDF2.PdfReader(file)
        #         for page in reader.pages:
        #             text = page.extract_text()
        #             if text:
        #                 cleaned_text = clean_text(text)
        #                 text_content += cleaned_text + "\n\n"
        # except Exception as e2:
        #     st.error(f"PyPDF2 also failed: {e2}")
    
    return text_content

def extract_topic_metadata_simple(text_chunk):
    """Simple rule-based topic extraction as fallback when Ollama is unavailable"""
    text_lower = text_chunk.lower()
    
    # Define keyword patterns for subjects
    subject_keywords = {
        "Operating System": ["process", "memory", "scheduling", "kernel", "thread", "deadlock", "virtual memory"],
        "Computer Network": ["tcp", "ip", "protocol", "network", "router", "switch", "osi", "lan", "wan"],
        "DBMS": ["database", "sql", "normalization", "transaction", "index", "query", "rdbms"],
        "OOPS": ["object", "class", "inheritance", "polymorphism", "encapsulation", "abstraction"],
        "System Design": ["scalability", "load balancer", "microservices", "cache", "database design"],
        "LLD": ["low level design", "class diagram", "uml", "design patterns"],
        "GitHub": ["git", "repository", "commit", "branch", "merge", "pull request"],
        "Linux": ["linux", "ubuntu", "centos", "kernel", "shell", "bash", "terminal"],
        "Aptitude": ["probability", "percentage", "ratio", "average", "time and work"]
    }
    
    # Find the best matching subject
    best_subject = "Operating System"
    max_matches = 0
    
    for subject, keywords in subject_keywords.items():
        matches = sum(1 for keyword in keywords if keyword in text_lower)
        if matches > max_matches:
            max_matches = matches
            best_subject = subject
    
    # Extract potential topics from the text
    sentences = re.split(r'[.!?]', text_chunk)
    potential_topics = []
    
    for sentence in sentences[:5]:  # Check first few sentences
        words = sentence.strip().split()
        if len(words) > 2 and len(words) < 8:
            potential_topics.append(sentence.strip())
    
    primary_topic = potential_topics[0] if potential_topics else "General Computing"
    
    return {
        "primary_topic": primary_topic[:50],  # Limit length
        "subtopics": potential_topics[1:4] if len(potential_topics) > 1 else ["basic concepts"],
        "subject_area": best_subject,
        "key_terms": list(set([word for word in text_lower.split()[:20] if len(word) > 4]))
    }

def extract_json_from_response(response):
    """Extract JSON from LLM response with better error handling"""
    response = response.strip()
    
    # Try to find JSON in the response
    json_pattern = r'\{.*\}'
    matches = re.findall(json_pattern, response, re.DOTALL)
    
    if matches:
        # Try each potential JSON match
        for match in matches:
            try:
                # Clean up the match
                cleaned = re.sub(r',\s*}', '}', match)  # Fix trailing commas
                cleaned = re.sub(r',\s*]', ']', cleaned)
                return json.loads(cleaned)
            except json.JSONDecodeError:
                continue
    
    # If no valid JSON found, return None
    return None

def extract_topic_metadata_ollama(text_chunk):
    """Extract topics using Ollama if available"""
    if not st.session_state.ollama_available:
        return extract_topic_metadata_simple(text_chunk)
    
    prompt = f"""
Analyze the following text and extract hierarchical topics and subtopics.
where topics and subtopics are in the form of 
example:{ "topics":["subtopics1", "subtopics2,..."]}

{
  "Operating System": [
    "Process Management",
    "Memory Management",
    "File Systems",
    "Deadlocks",
    "CPU Scheduling",
    "Threads and Concurrency",
    "Virtual Memory",
    "I/O Management",
    "System Calls",
    "Process Synchronization",
    "Inter-process Communication",
    "Storage Management"
  ],
  "Computer Network": [
    "TCP/IP",
    "OSI Model",
    "Network Topologies",
    "Routing Algorithms",
    "Network Security",
    "DNS",
    "HTTP/HTTPS",
    "Socket Programming",
    "Network Protocols",
    "Wireless Networks",
    "Network Layers",
    "Subnetting",
    "VPN"
  ],
  "DBMS": [
    "Database Normalization",
    "SQL Queries",
    "Indexing",
    "Transactions",
    "Concurrency Control",
    "Database Design",
    "Joins and Subqueries",
    "NoSQL Databases",
    "Database Security",
    "Query Optimization",
    "Data Warehousing",
    "Backup and Recovery"
  ],
  "OOPS": [
    "Object-Oriented Principles",
    "Classes and Objects",
    "Constructors and Destructors",
    "Method Overloading and Overriding",
    "Access Modifiers",
    "Interfaces and Abstract Classes",
    "Exception Handling",
    "Composition vs Inheritance",
    "Design Principles",
    "Generic Programming"
  ],
  "System Design": [
    "Design Patterns",
    "Scalability",
    "Load Balancing",
    "Caching Strategies",
    "Database Sharding",
    "Microservices Architecture",
    "API Design",
    "Message Queues",
    "CDN",
    "CAP Theorem",
    "System Architecture Patterns",
    "Performance Optimization"
  ],
  "LLD": [
    "Design Patterns",
    "UML Diagrams",
    "Object Modeling",
    "Code Organization",
    "Interface Design",
    "Design Principles",
    "Class Relationships",
    "Modular Design",
    "Testability",
    "Refactoring"
  ],
  "Git": [
    "Version Control",
    "Branching and Merging",
    "Git Workflows",
    "Rebasing",
    "Stashing",
    "Cherry-picking",
    "Git Hooks",
    "Conflict Resolution",
    "Git Commands",
    "Repository Management"
  ],
  "Linux": [
    "Linux Commands",
    "File Permissions",
    "Process Management",
    "Shell Scripting",
    "System Administration",
    "Package Management",
    "Networking Commands",
    "Text Processing",
    "Cron Jobs",
    "System Monitoring",
    "User Management"
  ],
  "Aptitude": [
    "Quantitative Aptitude",
    "Logical Reasoning",
    "Verbal Ability",
    "Data Interpretation",
    "Probability and Statistics",
    "Permutations and Combinations",
    "Time and Work",
    "Profit and Loss",
    "Ages and Ratios",
    "Coding Problems",
    "Pattern Recognition"
  ]
}
Return ONLY as valid JSON in the format:
{{
    "topic": "main topic name",
    "subtopics": ["subtopic1", "subtopic2", "subtopic3"]
}}

Text:
{text_chunk[:1500]}
"""
    
    try:
        # Use ONLY langchain_community.llms.Ollama
        from langchain_community.llms import Ollama
        llm = Ollama(model="llama3.2")

        response = llm.invoke(prompt)

        # Try to extract JSON
        metadata = extract_json_from_response(response)

        if metadata:
            # Validate subject area
            if metadata.get("subject_area") not in SUBJECTS:
                metadata["subject_area"] = "Operating System"
            return metadata

        else:
            st.warning("Could not parse JSON from Ollama response, using fallback")
            return extract_topic_metadata_simple(text_chunk)

    except Exception as e:
        st.warning(f"Ollama metadata extraction failed, using fallback: {e}")
        return extract_topic_metadata_simple(text_chunk)
    

def create_simple_vector_index(chunks_with_metadata):
    """Create a simple vector index using TF-IDF when FAISS is not available"""
    from sklearn.feature_extraction.text import TfidfVectorizer
    from sklearn.metrics.pairwise import cosine_similarity
    import pandas as pd
    
    texts = [chunk["text"] for chunk in chunks_with_metadata]
    metadatas = [chunk["metadata"] for chunk in chunks_with_metadata]
    
    # Create TF-IDF vectors
    vectorizer = TfidfVectorizer(max_features=1000, stop_words='english')
    tfidf_matrix = vectorizer.fit_transform(texts)
    
    return {
        "vectorizer": vectorizer,
        "tfidf_matrix": tfidf_matrix,
        "texts": texts,
        "metadatas": metadatas
    }

def create_faiss_index(chunks_with_metadata):
    """Create FAISS vector store from text chunks with metadata"""
    try:
        from langchain_community.embeddings import OllamaEmbeddings
        from langchain_community.vectorstores import FAISS
        
        texts = [chunk["text"] for chunk in chunks_with_metadata]
        metadatas = [
            {
                "primary_topic": chunk["metadata"]["primary_topic"],
                "subtopics": json.dumps(chunk["metadata"]["subtopics"]),
                "subject_area": chunk["metadata"]["subject_area"],
                "key_terms": json.dumps(chunk["metadata"]["key_terms"]),
                "chunk_id": i
            }
            for i, chunk in enumerate(chunks_with_metadata)
        ]
        
        embeddings = OllamaEmbeddings(model="llama3.2")
        vectorstore = FAISS.from_texts(texts, embeddings, metadatas=metadatas)
        return vectorstore
    except Exception as e:
        st.warning(f"FAISS index creation failed, using simple TF-IDF: {e}")
        return create_simple_vector_index(chunks_with_metadata)

def search_related_chunks(query, k=5):
    """Search for related chunks using available method"""
    if st.session_state.vector_store is None:
        st.error("No content processed yet. Please upload and process a PDF first.")
        return []
    
    try:
        # Check if it's a FAISS store or simple index
        if hasattr(st.session_state.vector_store, 'similarity_search'):
            # FAISS search
            docs = st.session_state.vector_store.similarity_search(query, k=k)
            return docs
        else:
            # Simple TF-IDF search
            from sklearn.metrics.pairwise import cosine_similarity
            import numpy as np
            
            simple_index = st.session_state.vector_store
            query_vec = simple_index["vectorizer"].transform([query])
            similarities = cosine_similarity(query_vec, simple_index["tfidf_matrix"]).flatten()
            
            # Get top k similar chunks
            top_indices = np.argsort(similarities)[-k:][::-1]
            
            # Create mock document objects
            class SimpleDoc:
                def __init__(self, page_content, metadata):
                    self.page_content = page_content
                    self.metadata = metadata
            
            results = []
            for idx in top_indices:
                if similarities[idx] > 0.1:  # Minimum similarity threshold
                    doc = SimpleDoc(
                        simple_index["texts"][idx],
                        simple_index["metadatas"][idx]
                    )
                    results.append(doc)
            
            return results
    except Exception as e:
        st.error(f"Search error: {e}")
        return []

def generate_content_fallback(related_chunks, content_type="flashcards", num_items=5):
    """Generate simple content when Ollama is unavailable"""
    combined_text = "\n".join([doc.page_content[:200] for doc in related_chunks[:3]])  # Use first 3 chunks
    
    if content_type == "flashcards":
        # Create simple flashcards based on sentences
        sentences = re.split(r'[.!?]', combined_text)
        flashcards = []
        
        for i, sentence in enumerate(sentences[:num_items]):
            if len(sentence.strip()) > 20:
                words = sentence.strip().split()
                if len(words) > 5:
                    # Simple question-answer format
                    front = f"What is {words[0]} {words[1]}?"
                    back = sentence.strip()
                    
                    flashcards.append({
                        "front": front,
                        "back": back,
                        "topic": related_chunks[0].metadata.get('primary_topic', 'General'),
                        "difficulty": "medium"
                    })
        
        return flashcards
    
    else:  # questions
        # Create simple questions
        sentences = re.split(r'[.!?]', combined_text)
        questions = []
        
        question_types = ["mcq", "short", "long"]
        
        for i, sentence in enumerate(sentences[:num_items]):
            if len(sentence.strip()) > 20:
                q_type = question_types[i % len(question_types)]
                
                question = {
                    "question": f"Explain: {sentence.strip()}",
                    "type": q_type,
                    "topic": related_chunks[0].metadata.get('primary_topic', 'General'),
                    "difficulty": "medium",
                    "cognitive_level": "understanding"
                }
                
                if q_type == "mcq":
                    question["options"] = [
                        "Correct answer related to the topic",
                        "Plausible but incorrect option",
                        "Another incorrect option", 
                        "Clearly wrong option"
                    ]
                
                questions.append(question)
        
        return questions

def generate_flashcards(related_chunks, num_flashcards=5):
    """Generate flashcards from related chunks"""
    if not st.session_state.ollama_available:
        return generate_content_fallback(related_chunks, "flashcards", num_flashcards)
    
    combined_text = "\n\n".join([doc.page_content for doc in related_chunks])
    
    prompt = f"""
Based on the following content about {related_chunks[0].metadata.get('primary_topic', 'the topic')}, 
create {num_flashcards} high-quality flashcards.

For EACH flashcard, provide:
- front: The question or prompt (clear and concise)
- back: The answer or explanation (comprehensive but focused)
- topic: The specific subtopic
- difficulty: easy/medium/hard

Ensure flashcards cover different aspects and are educational valuable.

Content:
{combined_text[:3000]}

Return ONLY as valid JSON array.
"""
    
    try:
        # Try new langchain-ollama package first
        try:
            from langchain_ollama import OllamaLLM
            llm = OllamaLLM(model="llama3.2")
        except ImportError:
            # Fall back to deprecated version
            from langchain_community.llms import Ollama
            llm = Ollama(model="llama3.2")
        
        response = llm.invoke(prompt)
        
        # Extract JSON from response
        flashcards = extract_json_from_response(response)
        
        if flashcards:
            return flashcards
        else:
            st.error("Failed to parse flashcards JSON, using fallback")
            return generate_content_fallback(related_chunks, "flashcards", num_flashcards)
            
    except Exception as e:
        st.error(f"Failed to generate flashcards with Ollama, using fallback: {e}")
        return generate_content_fallback(related_chunks, "flashcards", num_flashcards)

def generate_questions(related_chunks, question_type="mixed", num_questions=5):
    """Generate questions from related chunks based on type"""
    if not st.session_state.ollama_available:
        return generate_content_fallback(related_chunks, "questions", num_questions)
    
    combined_text = "\n\n".join([doc.page_content for doc in related_chunks])
    
    type_instructions = {
        "mcq": "Generate ONLY Multiple Choice Questions (MCQs) with 4 options each",
        "short": "Generate ONLY Short Answer questions", 
        "long": "Generate ONLY Long Answer questions",
        "mixed": "Generate a mix of: 40% MCQs, 30% Short Answer, 30% Long Answer"
    }
    
    prompt = f"""
Based on the following content about {related_chunks[0].metadata.get('primary_topic', 'the topic')}, 
create {num_questions} high-quality questions.

{type_instructions[question_type]}

For EACH question provide:
- question: The question text
- type: mcq/short/long
- topic: Specific subtopic
- difficulty: easy/medium/hard
- cognitive_level: remembering/understanding/applying/analyzing/evaluating/creating
- options: [array of 4 options] - ONLY for MCQ type

Content:
{combined_text[:3000]}

Return ONLY as valid JSON array.
"""
    
    try:
        # Try new langchain-ollama package first
        try:
            from langchain_ollama import OllamaLLM
            llm = OllamaLLM(model="llama3.2")
        except ImportError:
            # Fall back to deprecated version
            from langchain_community.llms import Ollama
            llm = Ollama(model="llama3.2")
        
        response = llm.invoke(prompt)
        
        # Extract JSON from response
        questions = extract_json_from_response(response)
        
        if questions:
            return questions
        else:
            st.error("Failed to parse questions JSON, using fallback")
            return generate_content_fallback(related_chunks, "questions", num_questions)
            
    except Exception as e:
        st.error(f"Failed to generate questions with Ollama, using fallback: {e}")
        return generate_content_fallback(related_chunks, "questions", num_questions)

def save_flashcards_docx(flashcards, filename):
    """Save flashcards to Word document"""
    doc = Document()
    doc.add_heading('Generated Flashcards', 0)
    
    for i, card in enumerate(flashcards):
        doc.add_paragraph(f"Flashcard {i+1}:")
        doc.add_paragraph(f"Front: {card['front']}")
        doc.add_paragraph(f"Back: {card['back']}")
        doc.add_paragraph(f"Topic: {card['topic']} | Difficulty: {card['difficulty']}")
        doc.add_paragraph("\n")
    
    doc.save(filename)

def save_questions_docx(questions, filename):
    """Save questions to Word document"""
    doc = Document()
    doc.add_heading('Generated Questions', 0)
    
    for i, q in enumerate(questions):
        doc.add_paragraph(f"Q{i+1}: {q['question']}")
        doc.add_paragraph(f"Type: {q['type']} | Topic: {q['topic']} | "
                         f"Difficulty: {q['difficulty']} | Cognitive: {q['cognitive_level']}")
        
        if q["type"] == "mcq" and "options" in q:
            doc.add_paragraph("Options:")
            for j, option in enumerate(q["options"], start=1):
                doc.add_paragraph(f"  {j}. {option}")
                
        doc.add_paragraph("\n")
    
    doc.save(filename)

def main():
    st.set_page_config(page_title="Smart Study Material Generator", layout="wide")
    
    # Check Ollama availability
    st.session_state.ollama_available = check_ollama_availability()
    
    st.title("🎯 Smart Study Material Generator")
    st.markdown("Upload materials, search by topic, and generate flashcards or questions on-demand")
    
    # Show Ollama status
    if st.session_state.ollama_available:
        st.sidebar.success("✅ Ollama is running")
    else:
        st.sidebar.warning("⚠️ Ollama not detected - using fallback methods")
        st.sidebar.info("To enable AI features, install and run Ollama:")
        st.sidebar.code("curl -fsSL https://ollama.ai/install.sh | sh\nollama pull llama3.2")
    
    # Sidebar for document processing
    with st.sidebar:
        st.header("📄 Document Processing")
        uploaded_file = st.file_uploader("Upload PDF", type=['pdf'])
        
        if uploaded_file:
            if st.button("Process Document"):
                with st.spinner("Processing document..."):
                    # Save uploaded file temporarily
                    with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp_file:
                        tmp_file.write(uploaded_file.getvalue())
                        pdf_path = tmp_file.name
                    
                    try:
                        # Extract text
                        extracted_text = extract_pdf_content(pdf_path)
                        
                        if extracted_text:
                            st.success(f"✅ Extracted {len(extracted_text)} characters")
                            
                            # Split into chunks
                            splitter = RecursiveCharacterTextSplitter(
                                chunk_size=1000,
                                chunk_overlap=200,
                                separators=["\n\n", "\n", ".", "!"]
                            )
                            chunks = splitter.split_text(extracted_text)
                            
                            # Extract metadata for each chunk
                            chunks_with_metadata = []
                            progress_bar = st.progress(0)
                            status_text = st.empty()
                            
                            for i, chunk in enumerate(chunks):
                                status_text.text(f"Processing chunk {i+1}/{len(chunks)}")
                                
                                if st.session_state.ollama_available:
                                    metadata = extract_topic_metadata_ollama(chunk)
                                else:
                                    metadata = extract_topic_metadata_simple(chunk)
                                
                                chunks_with_metadata.append({
                                    "text": chunk,
                                    "metadata": metadata
                                })
                                progress_bar.progress((i + 1) / len(chunks))
                            
                            status_text.text("Creating search index...")
                            
                            # Create search index
                            vector_store = create_faiss_index(chunks_with_metadata)
                            
                            # Store in session state
                            st.session_state.chunks = chunks_with_metadata
                            st.session_state.vector_store = vector_store
                            st.session_state.processed = True
                            
                            st.success(f"✅ Processed {len(chunks)} chunks!")
                            status_text.empty()
                        else:
                            st.error("No text could be extracted from the PDF")
                            
                    except Exception as e:
                        st.error(f"Processing failed: {e}")
                    finally:
                        if os.path.exists(pdf_path):
                            os.unlink(pdf_path)
        
        # Display processing status
        if st.session_state.processed:
            st.success(f"✅ {len(st.session_state.chunks)} chunks ready")
            
            # Show topics overview
            st.header("📊 Topics Found")
            all_topics = [chunk["metadata"]["primary_topic"] for chunk in st.session_state.chunks]
            topic_counts = pd.Series(all_topics).value_counts()
            for topic, count in topic_counts.head(8).items():
                st.write(f"• {topic}: {count} chunks")
    
    # Main content area
    if st.session_state.processed:
        st.header("🔍 Topic-Based Content Generation")
        
        # User input section
        col1, col2 = st.columns([2, 1])
        
        with col1:
            search_query = st.text_input(
                "Enter topic or concept to focus on:",
                placeholder="e.g., Inheritance, Process Scheduling, Database Normalization..."
            )
        
        with col2:
            num_chunks = st.slider("Number of relevant chunks", 3, 10, 5)
        
        # Content type selection
        content_type = st.radio(
            "What would you like to generate?",
            ["Flashcards", "Questions"],
            horizontal=True
        )
        
        # Advanced options
        with st.expander("Advanced Options"):
            if content_type == "Flashcards":
                num_items = st.slider("Number of flashcards", 3, 15, 8)
            else:
                col1, col2 = st.columns(2)
                with col1:
                    num_items = st.slider("Number of questions", 3, 15, 8)
                with col2:
                    question_type = st.selectbox(
                        "Question type",
                        ["mixed", "mcq", "short", "long"],
                        format_func=lambda x: {
                            "mixed": "Mixed Types", 
                            "mcq": "Multiple Choice",
                            "short": "Short Answer",
                            "long": "Long Answer"
                        }[x]
                    )
        
        # Generate button
        if st.button("Generate Content") and search_query:
            with st.spinner(f"Generating {content_type.lower()}..."):
                # Search for relevant chunks
                related_chunks = search_related_chunks(search_query, k=num_chunks)
                
                if related_chunks:
                    st.success(f"Found {len(related_chunks)} relevant chunks")
                    
                    # Display search results
                    with st.expander("View Relevant Content"):
                        for i, doc in enumerate(related_chunks):
                            st.write(f"**Chunk {i+1}** - Topic: {doc.metadata['primary_topic']}")
                            st.write(f"*Subject:* {doc.metadata['subject_area']}")
                            st.text(doc.page_content[:300] + "...")
                            st.divider()
                    
                    # Generate content based on type
                    if content_type == "Flashcards":
                        flashcards = generate_flashcards(related_chunks, num_flashcards=num_items)
                        
                        if flashcards:
                            st.header("📇 Generated Flashcards")
                            
                            # Display flashcards
                            for i, card in enumerate(flashcards):
                                with st.expander(f"Flashcard {i+1}: {card['topic']} ({card['difficulty']})"):
                                    col1, col2 = st.columns(2)
                                    with col1:
                                        st.subheader("Front")
                                        st.info(card['front'])
                                    with col2:
                                        st.subheader("Back") 
                                        st.success(card['back'])
                            
                            # Download option
                            if st.download_button(
                                "Download Flashcards (JSON)",
                                data=json.dumps(flashcards, indent=2),
                                file_name="flashcards.json",
                                mime="application/json"
                            ):
                                st.success("JSON downloaded!")
                    
                    else:  # Questions
                        questions = generate_questions(related_chunks, question_type, num_questions=num_items)
                        
                        if questions:
                            st.header("❓ Generated Questions")
                            
                            # Display questions
                            for i, q in enumerate(questions):
                                with st.expander(f"Q{i+1}: {q['topic']} ({q['difficulty']})"):
                                    st.write(f"**Question:** {q['question']}")
                                    st.write(f"**Type:** {q['type']} | **Cognitive Level:** {q['cognitive_level']}")
                                    
                                    if q["type"] == "mcq" and "options" in q:
                                        st.write("**Options:**")
                                        for j, option in enumerate(q["options"]):
                                            st.write(f"  {j+1}. {option}")
                            
                            # Download options
                            col1, col2 = st.columns(2)
                            with col1:
                                if st.download_button(
                                    "Download Questions (JSON)",
                                    data=json.dumps(questions, indent=2),
                                    file_name="questions.json",
                                    mime="application/json"
                                ):
                                    st.success("JSON downloaded!")
                
                else:
                    st.error("No relevant content found. Try a different search term.")
    
    else:
        # Welcome screen when no document is processed
        st.header("Welcome to Smart Study Material Generator!")
        
        col1, col2 = st.columns(2)
        
        with col1:
            st.subheader("🚀 How it works:")
            st.write("""
            1. **Upload** your study material (PDF)
            2. **Process** the document to extract topics
            3. **Search** for specific topics or concepts  
            4. **Generate** flashcards or questions
            """)
            
            if not st.session_state.ollama_available:
                st.warning("**Note:** Ollama is not running. The app will use fallback methods for content generation.")
        
        with col2:
            st.subheader("📚 Supported Subjects:")
            for subject in SUBJECTS:
                st.write(f"• {subject}")

if __name__ == "__main__":
    main()