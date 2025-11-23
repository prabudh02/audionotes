import streamlit as st
import os
import tempfile
import re
import json
import pandas as pd
from docx import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
import numpy as np
import requests
import pickle
from pathlib import Path

# Persistence paths (temp directory)
TMP_DIR = Path.cwd() / "smart_study_generator"
TMP_DIR.mkdir(parents=True, exist_ok=True)
FAISS_STORE_PATH = TMP_DIR / "faiss_store.pkl"
CHUNKS_PATH = TMP_DIR / "chunks.json"

# Subject categories
SUBJECTS = [
    "Operating System", "Computer Network", "DBMS", "OOPS",
    "System Design", "LLD", "Git", "Linux", "Aptitude"
]

def check_ollama_availability():
    """Check if Ollama is running and has at least one model"""
    try:
        response = requests.get("http://localhost:11434/api/tags", timeout=5)
        if response.status_code == 200:
            models = response.json().get('models', [])
            return len(models) > 0
        return False
    except Exception:
        return False

def clean_text(text):
    """Comprehensive text cleaning with word boundary reconstruction"""
    text = re.sub(r'(\b)e([A-Z])', r'\1\2', text)
    text = re.sub(r'\b(\d+);\b', r'\1', text)
    text = re.sub(r'([a-z])([A-Z])', r'\1 \2', text)
    text = re.sub(r'([A-Z])([A-Z][a-z])', r'\1 \2', text)
    text = re.sub(r'([0-9])([A-Za-z])', r'\1 \2', text)
    text = re.sub(r'([A-Za-z])([0-9])', r'\1 \2', text)
    text = re.sub(r'[^\w\s\.\-\+\*/\^\(\)\[\]\{\}=,;:&@#%$]', '', text)
    text = re.sub(r'\s+', ' ', text)
    text = re.sub(r'\s([\.\,\;:])', r'\1', text)
    text = re.sub(r'\bpage\s*\d+\b', '', text, flags=re.IGNORECASE)
    text = re.sub(r'(\w)-\s+(\w)', r'\1\2', text)
    replacements = {
        r'\boperatingsystem\b': 'operating system',
        r'\baddressspace\b': 'address space',
        r'\bmultitasking\b': 'multi-tasking',
        r'\bmainmemory\b': 'main memory',
        r'\bcpubound\b': 'CPU-bound',
        r'\biobound\b': 'I/O-bound',
        r'\bdatastructure\b': 'data structure',
        r'\balgorithm\b': 'algorithm',
    }
    for pattern, replacement in replacements.items():
        text = re.sub(pattern, replacement, text, flags=re.IGNORECASE)
    return text.strip()

def extract_pdf_content(pdf_path):
    """Extract text from PDF using OCR (pdf2image + pytesseract)"""
    text_content = ""
    try:
        import pdf2image
        import pytesseract
        import cv2
        images = pdf2image.convert_from_path(pdf_path)
        for i, image in enumerate(images):
            st.info(f"Processing page {i+1}/{len(images)}")
            gray = cv2.cvtColor(np.array(image), cv2.COLOR_RGB2GRAY)
            page_text = pytesseract.image_to_string(gray)
            cleaned_page_text = clean_text(page_text)
            text_content += cleaned_page_text + "\n\n"
    except Exception as e:
        # If OCR fails, surface the error so user can act
        raise RuntimeError(f"OCR extraction failed: {e}")
    return text_content


import json
import re

def extract_json_from_response(response):
    """Extracts JSON array or object from an LLM response safely."""
    response = response.strip()

    # Match either { ... } or [ ... ]
    json_pattern = r'(\{[\s\S]*\}|\[[\s\S]*\])'
    matches = re.findall(json_pattern, response)

    if not matches:
        return None

    for match in matches:
        cleaned = match

        # Fix trailing commas
        cleaned = re.sub(r',\s*([}\]])', r'\1', cleaned)

        # Try loading
        try:
            return json.loads(cleaned)
        except json.JSONDecodeError:
            continue

    return None


def extract_topic_metadata_ollama(text_chunk,i):
    """Extract topics using Ollama if available"""
    

    prompt = f"""
Analyze the following text and determine the most relevant topic and its subtopics.

Below is the list of allowed topics with their subtopics. Use ONLY these:

{{
  "Operating System": [
    "Process Management",
    "Memory Management",
    "File Systems",
    "Deadlocks",
    "CPU Scheduling",
    "Threads and Concurrency",
    "Virtual Memory",
    "I/O Management",
    "System Calls",
    "Process Synchronization",
    "Inter-process Communication",
    "Storage Management"
  ],
  "Computer Network": [
    "TCP/IP",
    "OSI Model",
    "Network Topologies",
    "Routing Algorithms",
    "Network Security",
    "DNS",
    "HTTP/HTTPS",
    "Socket Programming",
    "Network Protocols",
    "Wireless Networks",
    "Network Layers",
    "Subnetting",
    "VPN"
  ],
  "DBMS": [
    "Database Normalization",
    "SQL Queries",
    "Indexing",
    "Transactions",
    "Concurrency Control",
    "Database Design",
    "Joins and Subqueries",
    "NoSQL Databases",
    "Database Security",
    "Query Optimization",
    "Data Warehousing",
    "Backup and Recovery"
  ],
  "OOPS": [
    "Object-Oriented Principles",
    "Classes and Objects",
    "Constructors and Destructors",
    "Method Overloading and Overriding",
    "Access Modifiers",
    "Interfaces and Abstract Classes",
    "Exception Handling",
    "Composition vs Inheritance",
    "Design Principles",
    "Generic Programming"
  ],
  "System Design": [
    "Design Patterns",
    "Scalability",
    "Load Balancing",
    "Caching Strategies",
    "Database Sharding",
    "Microservices Architecture",
    "API Design",
    "Message Queues",
    "CDN",
    "CAP Theorem",
    "System Architecture Patterns",
    "Performance Optimization"
  ],
  "LLD": [
    "Design Patterns",
    "UML Diagrams",
    "Object Modeling",
    "Code Organization",
    "Interface Design",
    "Design Principles",
    "Class Relationships",
    "Modular Design",
    "Testability",
    "Refactoring"
  ],
  "Git": [
    "Version Control",
    "Branching and Merging",
    "Git Workflows",
    "Rebasing",
    "Stashing",
    "Cherry-picking",
    "Git Hooks",
    "Conflict Resolution",
    "Git Commands",
    "Repository Management"
  ],
  "Linux": [
    "Linux Commands",
    "File Permissions",
    "Process Management",
    "Shell Scripting",
    "System Administration",
    "Package Management",
    "Networking Commands",
    "Text Processing",
    "Cron Jobs",
    "System Monitoring",
    "User Management"
  ],
  "Aptitude": [
    "Quantitative Aptitude",
    "Logical Reasoning",
    "Verbal Ability",
    "Data Interpretation",
    "Probability and Statistics",
    "Permutations and Combinations",
    "Time and Work",
    "Profit and Loss",
    "Ages and Ratios",
    "Coding Problems",
    "Pattern Recognition"
  ]
}}
the above part is where the list of topics and subtopics is given using it as reference only.
Return ONLY valid JSON (no explanation) in this format:

{{
    "topic": "main topic name",
    "subtopics": ["sub1", "sub2", "sub3"]
}}

Text:
{text_chunk[:1500]}
"""

    try:
        from langchain_community.llms import Ollama
        llm = Ollama(model="llama3.2")

        response = llm.invoke(prompt)
        print("response:", response)
        metadata = extract_json_from_response(response)

        if metadata:
            return metadata
        else:
            st.warning("Could not parse JSON from Ollama response, using fallback:"+i)
            

    except Exception as e:
        st.warning(f"Ollama metadata extraction failed, using fallback: {e}")
        

  
        

  

def save_vectorstore_to_disk(vectorstore):
    """Persist the vectorstore object to disk (pickle)."""
    try:
        with open(FAISS_STORE_PATH, "wb") as f:
            pickle.dump(vectorstore, f)
    except Exception as e:
        raise RuntimeError(f"Failed to save FAISS store to disk: {e}")

def load_vectorstore_from_disk():
    """Load vectorstore from disk if exists"""
    if not FAISS_STORE_PATH.exists():
        return None
    try:
        with open(FAISS_STORE_PATH, "rb") as f:
            vs = pickle.load(f)
        return vs
    except Exception as e:
        raise RuntimeError(f"Failed to load FAISS store from disk: {e}")

def save_chunks_to_disk(chunks_with_metadata):
    try:
        with open(CHUNKS_PATH, "w", encoding="utf-8") as f:
            json.dump(chunks_with_metadata, f, ensure_ascii=False, indent=2)
    except Exception as e:
        raise RuntimeError(f"Failed to save chunks to disk: {e}")

def load_chunks_from_disk():
    if not CHUNKS_PATH.exists():
        return None
    try:
        with open(CHUNKS_PATH, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception as e:
        raise RuntimeError(f"Failed to load chunks from disk: {e}")

def create_faiss_index(chunks_with_metadata):
    """Create FAISS vector store from text chunks with metadata. Raises on failure."""
    try:
        from langchain_community.embeddings import OllamaEmbeddings
        from langchain.vectorstores import FAISS
    except Exception as e:
        raise RuntimeError(f"Required langchain_community modules for FAISS are not available: {e}")
    try:
        texts = [chunk["text"] for chunk in chunks_with_metadata]

        # FIX: ensure metadata is always a dict
        metadatas = []
        for i, chunk in enumerate(chunks_with_metadata):
            meta = chunk.get("metadata") or {}   # ensure dictionary, even if None

            metadatas.append({
                "topic": meta.get("topic"),
                "subtopics": json.dumps(meta.get("subtopics", [])),
                "chunk_id": i
            })

        embeddings = OllamaEmbeddings(model="llama3.2")

        vectorstore = FAISS.from_texts(
            texts=texts,
            embedding=embeddings,
            metadatas=metadatas
        )

        save_vectorstore_to_disk(vectorstore)
        return vectorstore

    except Exception as e:
        raise RuntimeError(f"FAISS index creation failed: {e}")

def search_related_chunks(query, k=10):
    """
    Search for related chunks. Expects FAISS vectorstore to be persisted or loaded.
    Raises/returns error if not available.
    """
    vectorstore = load_vectorstore_from_disk()
    if vectorstore is None:
        raise RuntimeError("No FAISS index found. Please process a document first to create a FAISS index.")
    try:
        # Use vectorstore's similarity_search (langchain FAISS wrapper)
        docs = vectorstore.similarity_search(query, k=k)
        # filtered = [
        # r for r in docs
        # if r.metadata.get("topic") == "query"
        # and "ACID" in r.metadata.get("subtopics", "")]

        return docs
    except Exception as e:
        raise RuntimeError(f"Search error: {e}")

def generate_flashcards(related_chunks, num_flashcards=5):
    """Generate flashcards using Ollama only. Raise on failure."""
    if not check_ollama_availability():
        raise RuntimeError("Ollama not available for flashcard generation.")
    combined_text = "\n\n".join([doc.page_content for doc in related_chunks])
    prompt = f"""
Based on the following content about  main topic:{related_chunks[0].metadata.get('topic', 'Core computer science topic')} and subtopic:{related_chunks[0].metadata.get('subtopics', 'Core computer science subtopic related to topic')}, using these
create {num_flashcards} high-quality detailed flashcards.

For EACH flashcard, provide a strict JSON object format with:
- front: question/prompt
- back: answer/explanation
- topic: the subtopic


Return ONLY as a valid JSON array.
Content:
{combined_text[:3000]}
"""
    try:
        from langchain_community.llms import Ollama
        llm = Ollama(model="llama3.2")
        response = llm.invoke(prompt)
        print("Flashcard response:", response)
        flashcards = extract_json_from_response(response)
        if not flashcards:
            raise RuntimeError("Failed to parse flashcards JSON from Ollama response.")
        return flashcards
    except Exception as e:
        raise RuntimeError(f"Flashcard generation failed: {e}")

def generate_questions(related_chunks, question_type="mixed", num_questions=5):
    """Generate questions using Ollama only. Raise on failure."""
    if not check_ollama_availability():
        raise RuntimeError("Ollama not available for question generation.")
    combined_text = "\n\n".join([doc.page_content for doc in related_chunks])
    type_instructions = {
        "mcq": "Generate ONLY Multiple Choice Questions (MCQs) with 4 options each",
        "short": "Generate ONLY Short Answer questions",
        "long": "Generate ONLY Long Answer questions",
        "mixed": "Generate a mix of: 40% MCQs, 30% Short Answer, 30% Long Answer"
    }
    prompt = f"""
Based on the following content about  main topic:{related_chunks[0].metadata.get('topic', 'Core computer science topic')} and subtopic:{related_chunks[0].metadata.get('subtopics', 'Core computer science subtopic related to topic')}, using these 
create {num_questions} high-quality questions.

{type_instructions[question_type]}

For EACH question provide a JSON object with:
- question
- type (mcq/short/long)
- topic
- difficulty
- cognitive_level
- options (array of 4) - ONLY for MCQ

Return ONLY as a valid JSON array.
Content:
{combined_text[:3000]}
"""
    try:
        from langchain_community.llms import Ollama
        llm = Ollama(model="llama3.2")
        response = llm.invoke(prompt)
        questions = extract_json_from_response(response)
        if not questions:
            raise RuntimeError("Failed to parse questions JSON from Ollama response.")
        return questions
    except Exception as e:
        raise RuntimeError(f"Question generation failed: {e}")

def save_flashcards_docx(flashcards, filename):
    doc = Document()
    doc.add_heading('Generated Flashcards', 0)
    for i, card in enumerate(flashcards):
        doc.add_paragraph(f"Flashcard {i+1}:")
        doc.add_paragraph(f"Front: {card['front']}")
        doc.add_paragraph(f"Back: {card['back']}")
        doc.add_paragraph(f"Topic: {card['topic']} ")
        doc.add_paragraph("\n")
    doc.save(filename)

def save_questions_docx(questions, filename):
    doc = Document()
    doc.add_heading('Generated Questions', 0)
    for i, q in enumerate(questions):
        doc.add_paragraph(f"Q{i+1}: {q['question']}")
        doc.add_paragraph(f"Type: {q['type']} | Topic: {q['topic']} | "
                          f"Difficulty: {q['difficulty']} | Cognitive: {q['cognitive_level']}")
        if q["type"] == "mcq" and "options" in q:
            doc.add_paragraph("Options:")
            for j, option in enumerate(q["options"], start=1):
                doc.add_paragraph(f"  {j}. {option}")
        doc.add_paragraph("\n")
    doc.save(filename)

def main():
    st.set_page_config(page_title="Smart Study Material Generator", layout="wide")
    st.title("🎯 Smart Study Material Generator (FAISS-only, Ollama-only)")
    st.markdown("This variant **does not** use `st.session_state`. It persists a local FAISS index and chunks to disk. Ollama is required for topic extraction and generation.")

    ollama_ok = check_ollama_availability()
    if ollama_ok:
        st.sidebar.success("✅ Ollama is running")
    else:
        st.sidebar.error("❌ Ollama not detected. This app requires Ollama for topic extraction and generation.")

    st.sidebar.header("📄 Document Processing")
    uploaded_file = st.sidebar.file_uploader("Upload PDF", type=['pdf'])

    if uploaded_file:
        if st.sidebar.button("Process Document"):
            with st.spinner("Processing document..."):
                # save upload to temp file
                with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_file:
                    tmp_file.write(uploaded_file.getvalue())
                    pdf_path = tmp_file.name
                try:
                    extracted_text = extract_pdf_content(pdf_path)
                    if not extracted_text or len(extracted_text.strip()) == 0:
                        st.error("No text could be extracted from the PDF.")
                        return

                    st.success(f"✅ Extracted {len(extracted_text)} characters")

                    splitter = RecursiveCharacterTextSplitter(
                        chunk_size=1000,
                        chunk_overlap=200,
                        separators=["\n\n", "\n", ".", "!"]
                    )
                    chunks = splitter.split_text(extracted_text)

                    chunks_with_metadata = []
                    progress_bar = st.progress(0)
                    status_text = st.empty()

                    for i, chunk in enumerate(chunks):
                        status_text.text(f"Processing chunk {i+1}/{len(chunks)}")

                        # MUST use Ollama for metadata - raise on failure
                        metadata = extract_topic_metadata_ollama(chunk,i)

                        if metadata:
                             chunks_with_metadata.append({
                            "text": chunk,
                            "metadata": metadata
                            })
                        progress_bar.progress((i + 1) / len(chunks))

                    status_text.text("Creating FAISS index...")
                    # create faiss index (will raise if something goes wrong)
                    vector_store = create_faiss_index(chunks_with_metadata)

                    # persist chunks metadata separately
                    save_chunks_to_disk(chunks_with_metadata)

                    st.success(f"✅ Processed {len(chunks_with_metadata)} chunks and created FAISS index.")
                    status_text.empty()

                except Exception as e:
                    st.error(f"Processing failed: {e}")
                finally:
                    if os.path.exists(pdf_path):
                        os.unlink(pdf_path)

    # If FAISS index exists, show topics
    chunks_on_disk = None
    if CHUNKS_PATH.exists():
        try:
            chunks_on_disk = load_chunks_from_disk()
            st.sidebar.success(f"✅ {len(chunks_on_disk)} chunks available on disk")
            st.sidebar.header("📊 Topics Found (top)")
            all_topics = [c["metadata"].get("primary_topic", "Unknown") for c in chunks_on_disk]
            topic_counts = pd.Series(all_topics).value_counts()
            for topic, count in topic_counts.head(8).items():
                st.sidebar.write(f"• {topic}: {count} chunks")
        except Exception as e:
            st.sidebar.error(f"Failed to read chunks from disk: {e}")

    # Main generation UI
    st.header("🔍 Topic-Based Content Generation (FAISS + Ollama required)")

    search_query = st.text_input(
        "Enter topic or concept to focus on:",
        placeholder="e.g., Inheritance, Process Scheduling, Database Normalization..."
    )
    num_chunks = st.slider("Number of relevant chunks", 3, 10, 5)

    content_type = st.radio(
        "What would you like to generate?",
        ["Flashcards", "Questions"],
        horizontal=True
    )

    if content_type == "Flashcards":
        num_items = st.slider("Number of flashcards", 3, 15, 8)
    else:
        col1, col2 = st.columns(2)
        with col1:
            num_items = st.slider("Number of questions", 3, 15, 8)
        with col2:
            question_type = st.selectbox(
                "Question type",
                ["mixed", "mcq", "short", "long"],
                format_func=lambda x: {
                    "mixed": "Mixed Types",
                    "mcq": "Multiple Choice",
                    "short": "Short Answer",
                    "long": "Long Answer"
                }[x]
            )

    if st.button("Generate Content") and search_query:
        try:
            related_chunks = search_related_chunks(search_query, k=num_chunks)
            if not related_chunks:
                st.error("No relevant chunks found (or FAISS returned empty).")
            else:
                st.success(f"Found {len(related_chunks)} relevant chunks")

                with st.expander("View Relevant Content"):
                    for i, doc in enumerate(related_chunks):
                        st.write(f"**Chunk {i+1}** - Topic: {doc.metadata.get('topic', 'Unknown')}")
                        st.write(f"*Subject:* {doc.metadata.get('subtopic', 'Unknown')}")
                        st.text(doc.page_content[:300] + "...")
                        st.divider()

                if content_type == "Flashcards":
                    flashcards = generate_flashcards(related_chunks, num_flashcards=num_items)
                    st.header("📇 Generated Flashcards")
                    for i, card in enumerate(flashcards):
                        with st.expander(f"Flashcard {i+1}: {card.get('topic','Unknown')} "):
                            col1, col2 = st.columns(2)
                            with col1:
                                st.subheader("Front")
                                st.info(card.get('front',''))
                            with col2:
                                st.subheader("Back")
                                st.success(card.get('back',''))
                    if flashcards:
                        st.download_button(
                            "Download Flashcards (JSON)",
                            data=json.dumps(flashcards, indent=2),
                            file_name="flashcards.json",
                            mime="application/json"
                        )
                else:
                    questions = generate_questions(related_chunks, question_type, num_questions=num_items)
                    st.header("❓ Generated Questions")
                    for i, q in enumerate(questions):
                        with st.expander(f"Q{i+1}: {q.get('topic','Unknown')} "):
                            st.write(f"**Question:** {q.get('question','')}")
                            st.write(f"**Type:** {q.get('type','')} | **Cognitive Level:** {q.get('cognitive_level','')}")
                            if q.get("type") == "mcq" and "options" in q:
                                st.write("**Options:**")
                                for j, option in enumerate(q["options"]):
                                    st.write(f"  {j+1}. {option}")
                    if questions:
                        st.download_button(
                            "Download Questions (JSON)",
                            data=json.dumps(questions, indent=2),
                            file_name="questions.json",
                            mime="application/json"
                        )
        except Exception as e:
            st.error(f"Generation failed: {e}")

    # Help / welcome
    if not CHUNKS_PATH.exists():
        st.header("Welcome to Smart Study Material Generator!")
        col1, col2 = st.columns(2)
        with col1:
            st.subheader("🚀 How it works:")
            st.write("""
            1. **Upload** your study material (PDF)
            2. **Process** the document to extract topics (Ollama required)
            3. **FAISS index** will be created and saved locally
            4. **Search** for topics and generate flashcards/questions (Ollama required)
            """)
            if not ollama_ok:
                st.warning("**Note:** Ollama is not running. This variant requires Ollama; processing/generation will fail without it.")
        with col2:
            st.subheader("📚 Supported Subjects:")
            for subject in SUBJECTS:
                st.write(f"• {subject}")

if __name__ == "__main__":
    main()
