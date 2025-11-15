import streamlit as st
import os
import tempfile
import re
import json
from PIL import Image
from docx import Document
import pandas as pd
import sys

# Add compatibility check at the start
try:
    import numpy as np
    print(f"NumPy version: {np.__version__}")
except ImportError as e:
    st.error(f"NumPy import error: {e}")

# Initialize session state
if 'processed' not in st.session_state:
    st.session_state.processed = False
if 'questions' not in st.session_state:
    st.session_state.questions = []
if 'vector_store' not in st.session_state:
    st.session_state.vector_store = None

# Subject categories
SUBJECTS = [
    "Operating System", "Computer Network", "DBMS", "OOPS", 
    "System Design", "LLD", "GitHub", "Linux", "Aptitude"
]

def clean_text(text):
    """Comprehensive text cleaning with word boundary reconstruction"""
    # Step 1: Fix specific OCR artifacts
    text = re.sub(r'(\b)e([A-Z])', r'\1\2', text)
    text = re.sub(r'\b(\d+);\b', r'\1', text)
    
    # Step 2: Reconstruct word boundaries
    text = re.sub(r'([a-z])([A-Z])', r'\1 \2', text)
    text = re.sub(r'([A-Z])([A-Z][a-z])', r'\1 \2', text)
    text = re.sub(r'([0-9])([A-Za-z])', r'\1 \2', text)
    text = re.sub(r'([A-Za-z])([0-9])', r'\1 \2', text)
    
    # Step 3: Standard cleaning
    text = re.sub(r'[^\w\s\.\-\+\*/\^\(\)\[\]\{\}=,;:&@#%$]', '', text)
    text = re.sub(r'\s+', ' ', text)
    text = re.sub(r'\s([\.\,\;:])', r'\1', text)
    text = re.sub(r'\bpage\s*\d+\b', '', text, flags=re.IGNORECASE)
    text = re.sub(r'(\w)-\s+(\w)', r'\1\2', text)
    
    # Step 4: Fix common OCR errors
    replacements = {
        r'\boperatingsystem\b': 'operating system',
        r'\baddressspace\b': 'address space',
        r'\bmultitasking\b': 'multi-tasking',
        r'\bmainmemory\b': 'main memory',
        r'\bcpubound\b': 'CPU-bound',
        r'\biobound\b': 'I/O-bound',
    }
    for pattern, replacement in replacements.items():
        text = re.sub(pattern, replacement, text, flags=re.IGNORECASE)
    
    return text.strip()

def extract_pdf_content(pdf_path):
    """Extract text from PDF using alternative methods"""
    text_content = ""
    
    try:
        # Try using pdf2image + pytesseract
        import pdf2image
        import pytesseract
        import cv2
        import numpy as np
        
        images = pdf2image.convert_from_path(pdf_path)
        
        for i, image in enumerate(images):
            st.info(f"Processing page {i+1}/{len(images)}")
            
            # Convert to grayscale
            gray = cv2.cvtColor(np.array(image), cv2.COLOR_RGB2GRAY)
            
            # Extract text using OCR
            page_text = pytesseract.image_to_string(gray)
            cleaned_page_text = clean_text(page_text)
            text_content += cleaned_page_text + "\n\n"
            
    except Exception as e:
        st.warning(f"OCR method failed: {e}. Trying PyPDF2 as fallback...")
        # Fallback to PyPDF2
        try:
            import PyPDF2
            with open(pdf_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                for page in reader.pages:
                    text = page.extract_text()
                    if text:
                        cleaned_text = clean_text(text)
                        text_content += cleaned_text + "\n\n"
        except Exception as e2:
            st.error(f"PyPDF2 also failed: {e2}")
    
    return text_content

def detect_subject_topic(text_chunk):
    """Use LLM to detect subject and topic with metadata"""
    prompt = f"""
Analyze the following text and determine:
1. Primary subject (choose from: {", ".join(SUBJECTS)})
2. Specific topic (2-5 words)
3. Key concepts (3-5 main concepts)

Return ONLY as valid JSON:
{{
    "subject": "subject_name",
    "topic": "specific_topic",
    "key_concepts": ["concept1", "concept2", "concept3"]
}}

Text:
{text_chunk[:1000]}
"""
    
    try:
        from langchain_community.llms import Ollama
        llm = Ollama(model="llama3.2")
        response = llm.invoke(prompt)
        
        # Clean the response
        response = response.strip()
        if "```json" in response:
            response = response.split("```json")[1].split("```")[0]
        elif "```" in response:
            response = response.split("```")[1].split("```")[0]
            
        metadata = json.loads(response)
        
        # Validate subject
        if metadata["subject"] not in SUBJECTS:
            metadata["subject"] = "Operating System"
            
        return metadata
    except Exception as e:
        st.warning(f"Failed to extract metadata: {e}")
        return {
            "subject": "Operating System",
            "topic": "General Computing",
            "key_concepts": ["computing", "technology", "systems"]
        }

def generate_questions(chunk, metadata):
    """Generate questions based on text chunk and metadata"""
    prompt = f"""
Generate diverse questions from this text chunk about {metadata['topic']} in {metadata['subject']}.

Key concepts: {', '.join(metadata['key_concepts'])}

Include:
1. 1 MCQ with 4 options
2. 1 Short answer question  
3. 1 Long answer question
4. 1 True/False question

For EACH question provide:
- question text
- topic (1-3 words)
- type (mcq/short/long/true_false)
- difficulty (easy/medium/hard)
- cognitive level (remembering/understanding/applying/evaluating)
- subject: "{metadata['subject']}"
- For MCQ ONLY: "options" array with exactly 4 choices
- For True/False: "options" array with ["True", "False"]

Output ONLY as valid JSON array.

Text Chunk:
{chunk}
"""
    
    try:
        from langchain_community.llms import Ollama
        llm = Ollama(model="llama3.2")
        response = llm.invoke(prompt)
        
        # Clean response
        response = response.strip()
        if "```json" in response:
            response = response.split("```json")[1].split("```")[0]
        elif "```" in response:
            response = response.split("```")[1].split("```")[0]
            
        questions = json.loads(response)
        
        # Add metadata to each question
        for q in questions:
            q.update({
                "subject": metadata["subject"],
                "main_topic": metadata["topic"],
                "key_concepts": metadata["key_concepts"]
            })
            
        return questions
    except Exception as e:
        st.warning(f"Failed to generate questions for chunk: {e}")
        return []

def create_faiss_index(questions):
    """Create FAISS vector store from questions"""
    try:
        from langchain_community.embeddings import OllamaEmbeddings
        from langchain_community.vectorstores import FAISS
        
        texts = [q["question"] for q in questions]
        metadatas = [
            {
                "topic": q["topic"],
                "type": q["type"],
                "difficulty": q["difficulty"],
                "cognitive_level": q["cognitive_level"],
                "subject": q["subject"],
                "main_topic": q["main_topic"],
                **({"options": json.dumps(q["options"])} if "options" in q else {})
            }
            for q in questions
        ]
        
        embeddings = OllamaEmbeddings(model="llama3.2")
        vectorstore = FAISS.from_texts(texts, embeddings, metadatas=metadatas)
        return vectorstore
    except Exception as e:
        st.error(f"Error creating FAISS index: {e}")
        return None

def save_questions_docx(questions, filename):
    """Save questions to Word document"""
    doc = Document()
    doc.add_heading('Generated Question Bank', 0)
    
    for i, q in enumerate(questions):
        doc.add_paragraph(f"Q{i+1}: {q['question']}")
        doc.add_paragraph(f"Subject: {q['subject']} | Topic: {q['topic']} | "
                         f"Type: {q['type']} | Difficulty: {q['difficulty']} | "
                         f"Cognitive: {q['cognitive_level']}")
        
        if "options" in q:
            doc.add_paragraph("Options:")
            for j, option in enumerate(q["options"], start=1):
                doc.add_paragraph(f"  {j}. {option}")
                
        doc.add_paragraph("\n")
    
    doc.save(filename)

def main():
    st.set_page_config(page_title="RAG Study Material Processor", layout="wide")
    
    st.title("📚 RAG Study Material Processor")
    st.markdown("Extract text from materials and generate flashcards/questions with metadata")
    
    # Check for required dependencies
    with st.sidebar:
        st.header("Dependency Check")
        
        missing_deps = []
        try:
            import numpy as np
            st.success("✓ NumPy")
        except:
            missing_deps.append("numpy")
            st.error("✗ NumPy")
            
        try:
            import pdf2image
            st.success("✓ pdf2image")
        except:
            missing_deps.append("pdf2image")
            st.error("✗ pdf2image")
            
        try:
            import pytesseract
            st.success("✓ pytesseract")
        except:
            missing_deps.append("pytesseract")
            st.error("✗ pytesseract")
            
        try:
            import langchain_community
            st.success("✓ langchain")
        except:
            missing_deps.append("langchain-community")
            st.error("✗ langchain")
        
        if missing_deps:
            st.error(f"Missing dependencies: {', '.join(missing_deps)}")
            st.info("Run: pip install " + " ".join(missing_deps))
        
        st.header("Configuration")
        uploaded_file = st.file_uploader("Upload PDF", type=['pdf'])
        process_btn = st.button("Process Material")
        
        st.header("Subjects")
        selected_subjects = st.multiselect(
            "Focus subjects:",
            SUBJECTS,
            default=["Operating System", "Computer Network", "DBMS"]
        )

    # Main content area
    if uploaded_file and process_btn:
        with st.spinner("Processing your document..."):
            # Save uploaded file temporarily
            with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp_file:
                tmp_file.write(uploaded_file.getvalue())
                pdf_path = tmp_file.name
            
            try:
                # Step 1: OCR and Text Extraction
                st.subheader("📄 Text Extraction")
                extracted_text = extract_pdf_content(pdf_path)
                
                if extracted_text:
                    st.success(f"Extracted {len(extracted_text)} characters of text")
                    
                    with st.expander("View Extracted Text"):
                        st.text_area("Cleaned Text", extracted_text, height=200)
                    
                    # Step 2: Text Chunking
                    st.subheader("🔪 Text Chunking")
                    from langchain_text_splitters import RecursiveCharacterTextSplitter
                    
                    splitter = RecursiveCharacterTextSplitter(
                        chunk_size=800,
                        chunk_overlap=100,
                        separators=["\n\n", "\n", ".", "!"]
                    )
                    chunks = splitter.split_text(extracted_text)
                    st.info(f"Created {len(chunks)} text chunks")
                    
                    # Step 3: Metadata Extraction and Question Generation
                    st.subheader("🏷️ Metadata Extraction & Question Generation")
                    all_questions = []
                    
                    progress_bar = st.progress(0)
                    for i, chunk in enumerate(chunks):
                        st.write(f"Processing chunk {i+1}/{len(chunks)}")
                        
                        # Extract metadata
                        metadata = detect_subject_topic(chunk)
                        
                        # Only process if subject matches selected subjects
                        if metadata["subject"] in selected_subjects:
                            # Generate questions
                            questions = generate_questions(chunk, metadata)
                            all_questions.extend(questions)
                            st.write(f"✅ Generated {len(questions)} questions")
                        else:
                            st.write(f"⏭️ Skipped - Subject: {metadata['subject']}")
                        
                        progress_bar.progress((i + 1) / len(chunks))
                    
                    # Save to session state
                    st.session_state.questions = all_questions
                    st.session_state.processed = True
                    
                    # Step 4: Create FAISS index
                    if all_questions:
                        st.subheader("🔍 Creating Search Index")
                        vector_store = create_faiss_index(all_questions)
                        st.session_state.vector_store = vector_store
                        st.success("FAISS index created successfully!")
                
                else:
                    st.error("No text could be extracted from the PDF")
                    
            except Exception as e:
                st.error(f"Processing failed: {e}")
            finally:
                # Clean up temporary file
                if os.path.exists(pdf_path):
                    os.unlink(pdf_path)
    
    # Display results if processing is complete
    if st.session_state.processed and st.session_state.questions:
        st.subheader("📊 Results Summary")
        
        # Statistics
        col1, col2, col3, col4 = st.columns(4)
        with col1:
            st.metric("Total Questions", len(st.session_state.questions))
        with col2:
            subjects_count = len(set(q['subject'] for q in st.session_state.questions))
            st.metric("Subjects Covered", subjects_count)
        with col3:
            mcq_count = len([q for q in st.session_state.questions if q['type'] == 'mcq'])
            st.metric("MCQ Questions", mcq_count)
        with col4:
            hard_count = len([q for q in st.session_state.questions if q['difficulty'] == 'hard'])
            st.metric("Hard Questions", hard_count)
        
        # Questions table
        st.subheader("📋 Generated Questions")
        df = pd.DataFrame(st.session_state.questions)
        
        # Display options for MCQ questions
        def format_options(row):
            if 'options' in row and row['options']:
                return "\n".join([f"{i+1}. {opt}" for i, opt in enumerate(row['options'])])
            return "N/A"
        
        if 'options' in df.columns:
            df['options_display'] = df.apply(format_options, axis=1)
        
        display_columns = ['question', 'subject', 'topic', 'type', 'difficulty']
        if 'options_display' in df.columns:
            display_columns.append('options_display')
            
        st.dataframe(df[display_columns], use_container_width=True)
        
        # Download options
        st.subheader("💾 Download Options")
        col1, col2, col3 = st.columns(3)
        
        with col1:
            if st.download_button(
                "Download Questions (JSON)",
                data=json.dumps(st.session_state.questions, indent=2),
                file_name="generated_questions.json",
                mime="application/json"
            ):
                st.success("JSON file downloaded!")
        
        with col2:
            # Create and download Word document
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_doc:
                save_questions_docx(st.session_state.questions, tmp_doc.name)
                with open(tmp_doc.name, "rb") as f:
                    if st.download_button(
                        "Download Questions (Word)",
                        data=f.read(),
                        file_name="question_bank.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    ):
                        st.success("Word document downloaded!")
                os.unlink(tmp_doc.name)
        
        with col3:
            # Export to CSV
            csv_data = df.to_csv(index=False)
            if st.download_button(
                "Download Questions (CSV)",
                data=csv_data,
                file_name="questions.csv",
                mime="text/csv"
            ):
                st.success("CSV file downloaded!")

if __name__ == "__main__":
    main()